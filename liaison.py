import streamlit as st
import pandas as pd
import numpy as np

# Check for critical dependencies
try:
    import matplotlib.pyplot as plt
    import seaborn as sns
    from scipy import stats
    from scipy.stats import shapiro, pearsonr, spearmanr, chi2_contingency, fisher_exact, bartlett, levene, f_oneway, kruskal, probplot
    from sklearn.linear_model import LinearRegression
    import plotly.express as px
    import plotly.graph_objects as go
    from statsmodels.stats.proportion import proportion_confint
    from io import BytesIO
    from docx import Document
    from docx.shared import Inches
    import base64
except ModuleNotFoundError as e:
    st.error(f"Erreur : Module manquant - {str(e)}. Veuillez installer les dépendances listées dans requirements.txt.")
    st.stop()

# --- Page Configuration ---
st.set_page_config(
    page_title="Analyse de Liaison Statistique",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- Custom CSS for Dark and Technological Theme ---
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Orbitron:wght@400;700&family=Roboto:wght@300;400;500&display=swap');

    /* Main app styling */
    body {
        font-family: 'Roboto', sans-serif;
        color: #e0e6ed !important;
    }
    .stApp {
        background: linear-gradient(135deg, #0c0c0c 0%, #1a1a2e 50%, #16213e 100%);
    }

    /* Sidebar styling */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #0f0f23 0%, #1a1a2e 100%) !important;
        border-right: 2px solid #00ffff;
        box-shadow: 5px 0 15px rgba(0, 255, 255, 0.1);
    }
    [data-testid="stSidebar"] .stFileUploader label,
    [data-testid="stSidebar"] .stSelectbox label,
    [data-testid="stSidebar"] .stSlider label,
    [data-testid="stSidebar"] .stButton button,
    [data-testid="stSidebar"] .stRadio label {
        color: #b0c4de !important;
        font-family: 'Roboto', sans-serif;
    }
    [data-testid="stSidebar"] .stButton button {
        background-color: #007bff;
        color: white;
        font-weight: bold;
        border-radius: 5px;
    }

    /* Main title */
    h1 {
        color: #00ffff;
        text-align: center;
        font-family: 'Orbitron', monospace;
        text-shadow: 0 0 10px rgba(0, 255, 255, 0.5);
        padding-top: 1.5rem;
    }

    /* Sub-headers */
    h2, h3, h4 {
        color: #00ffff;
        font-family: 'Orbitron', sans-serif;
        text-shadow: 0 0 8px rgba(0, 255, 255, 0.4);
    }

    /* Selectbox and radio styling */
    .stSelectbox, .stRadio {
        background-color: rgba(15, 15, 35, 0.8) !important;
        border: 1px solid #00ffff;
        border-radius: 5px;
        padding: 10px;
    }

    /* Tab styling */
    .stTabs [data-baseweb="tab-list"] {
        background: linear-gradient(90deg, #0f3460 0%, #16537e 100%);
        border-bottom: 2px solid #00ffff;
        border-radius: 10px 10px 0 0;
    }
    .stTabs [data-baseweb="tab"] {
        color: #b0c4de !important;
        font-family: 'Roboto', sans-serif;
        font-weight: bold;
    }
    .stTabs [aria-selected="true"] {
        background: #00ffff !important;
        color: #0f0f23 !important;
        text-shadow: none;
    }

    /* Expander styling */
    .st-expander {
        border: 1px solid #00ffff;
        border-radius: 10px;
        background: rgba(15, 15, 35, 0.8);
    }
    .st-expander header {
        font-size: 1.2rem;
        color: #00ffff;
        font-family: 'Orbitron', monospace;
    }

    /* Metric styling */
    [data-testid="stMetric"] {
        background: linear-gradient(135deg, rgba(15, 15, 35, 0.9) 0%, rgba(26, 26, 46, 0.9) 100%);
        border: 1px solid #00ffff;
        border-radius: 10px;
        padding: 15px;
        text-align: center;
    }

    /* DataFrame styling */
    .stDataFrame {
        background: rgba(15, 15, 35, 0.8) !important;
        border: 1px solid rgba(0, 255, 255, 0.3);
    }

    /* Uploader styling in sidebar */
    [data-testid="stFileUploader"] {
        border: 2px dashed #00ffff;
        background-color: #1a1a2e;
        padding: 20px;
        border-radius: 10px;
    }

    /* Alert styling */
    .stAlert {
        border-radius: 0.5rem;
        background: rgba(15, 15, 35, 0.8);
        border: 1px solid #00ffff;
        color: #e0e6ed !important;
    }

    /* Value box styling */
    .value-box {
        padding: 15px;
        border-radius: 5px;
        margin-bottom: 15px;
        text-align: center;
        color: #e0e6ed;
        background: linear-gradient(135deg, rgba(15, 15, 35, 0.9) 0%, rgba(26, 26, 46, 0.9) 100%);
        border: 1px solid #00ffff;
    }
    .value-box.primary {
        background: linear-gradient(135deg, #0f3460 0%, #16537e 100%);
    }
    .value-box.success {
        background: linear-gradient(135deg, #00ff88 0%, #00cc66 100%);
    }
    .value-box.warning {
        background: linear-gradient(135deg, #ffaa00 0%, #cc8800 100%);
    }
    .value-box.danger {
        background: linear-gradient(135deg, #ff4444 0%, #cc3333 100%);
    }
    .value-box.info {
        background: linear-gradient(135deg, #00cccc 0%, #009999 100%);
    }

    /* Uploaded file styling */
    .uploaded-file {
        color: #00ff88 !important;
        font-weight: bold;
    }

    /* Footer styling */
    .footer {
        font-size: 0.9rem;
        color: #b0c4de;
        text-align: center;
        padding: 1rem;
        margin-top: 2rem;
        background: linear-gradient(135deg, rgba(15, 15, 35, 0.9) 0%, rgba(26, 26, 46, 0.9) 100%);
        border: 1px solid #00ffff;
        border-radius: 5px;
    }

    /* Author info box */
    .author-info {
        background: linear-gradient(135deg, rgba(15, 15, 35, 0.9) 0%, rgba(26, 26, 46, 0.9) 100%);
        border: 2px solid #00ffff;
        border-radius: 15px;
        padding: 20px;
        margin-top: 20px;
        box-shadow: 0 10px 30px rgba(0, 255, 255, 0.2);
    }
</style>
""", unsafe_allow_html=True)

# --- Plotting Theme Configuration ---
plt.style.use('dark_background')
plt.rcParams.update({
    'figure.facecolor': '#0c0c0c',
    'axes.facecolor': '#1a1a2e',
    'axes.edgecolor': '#00ffff',
    'axes.labelcolor': '#e0e6ed',
    'xtick.color': '#e0e6ed',
    'ytick.color': '#e0e6ed',
    'grid.color': '#333333',
    'text.color': '#e0e6ed',
    'legend.facecolor': '#0f0f23',
    'legend.edgecolor': '#333333'
})

# Plotly Dark Theme
plotly_dark_template = go.layout.Template(
    layout=dict(
        plot_bgcolor='#1a1a2e',
        paper_bgcolor='#0c0c0c',
        font_color='#e0e6ed',
        xaxis=dict(gridcolor='#333333', linecolor='#e0e6ed'),
        yaxis=dict(gridcolor='#333333', linecolor='#e0e6ed'),
        title_font_color='#00ffff',
        xaxis_title_font_color='#00ffff',
        yaxis_title_font_color='#00ffff',
        legend=dict(bgcolor='rgba(15,15,35,0.8)', bordercolor='#333333')
    )
)

# --- Application Title ---
st.title("📊 Analyse de Liaison Statistique")
st.markdown("<p style='text-align: center; color: #b0c4de; font-family: Roboto, sans-serif;'>Une plateforme pour analyser les relations entre variables quantitatives et qualitatives avec des tests statistiques robustes.</p>", unsafe_allow_html=True)

# --- Demo Dataset ---
def load_demo_data():
    from sklearn.datasets import load_iris
    iris = load_iris()
    df = pd.DataFrame(data=iris.data, columns=iris.feature_names)
    df['species'] = pd.Categorical([iris.target_names[i] for i in iris.target])
    df['quality'] = pd.Categorical(np.random.choice(['Low', 'Medium', 'High'], size=len(df)))
    return df

# --- Value Box Function ---
def value_box(title, value, color="primary", icon=None):
    icon_html = f"<span>{icon}</span> " if icon else ""
    st.markdown(f"""
        <div class="value-box {color}">
            <h4>{value}</h4>
            <p>{icon_html}{title}</p>
        </div>
        """, unsafe_allow_html=True)

# --- Report Generation Functions ---
def generate_quant_quant_report(df, var1, var2, corr, p_value, coef_symbol, normal_var1, normal_var2):
    doc = Document()
    doc.add_heading('Rapport d\'Analyse de Liaison : Quantitative vs Quantitative', 0)
    doc.add_heading(f'Variables analysées : {var1} vs {var2}', level=1)
    
    doc.add_heading('Résultats des Tests de Normalité', level=2)
    doc.add_paragraph(f'{var1} : {"Normale" if normal_var1 else "Non normale"} (p-valeur = {shapiro(df[var1])[1]:.4f})')
    doc.add_paragraph(f'{var2} : {"Normale" if normal_var2 else "Non normale"} (p-valeur = {shapiro(df[var2])[1]:.4f})')
    
    doc.add_heading('Résultats du Test de Corrélation', level=2)
    doc.add_paragraph(f'Méthode utilisée : {"Pearson" if normal_var1 and normal_var2 else "Spearman"}')
    doc.add_paragraph(f'Coefficient ({coef_symbol}) : {corr:.4f}')
    doc.add_paragraph(f'p-valeur : {p_value:.4f}')
    doc.add_paragraph(f'Conclusion : {"Pas d\'association significative" if p_value > 0.05 else "Association significative"}')
    
    output = BytesIO()
    doc.save(output)
    return output.getvalue()

def generate_qual_qual_report(input_tab, method, stat_value, dof_val, p_val):
    doc = Document()
    doc.add_heading('Rapport d\'Analyse de Liaison : Qualitative vs Qualitative', 0)
    doc.add_heading('Tableau de Contingence', level=1)
    
    table = doc.add_table(rows=input_tab.shape[0]+1, cols=input_tab.shape[1]+1)
    table.style = 'Table Grid'
    for j, col in enumerate(input_tab.columns):
        table.cell(0, j+1).text = str(col)
    for i, idx in enumerate(input_tab.index):
        table.cell(i+1, 0).text = str(idx)
        for j, val in enumerate(input_tab.iloc[i]):
            table.cell(i+1, j+1).text = str(val)
    
    doc.add_heading('Résultats du Test', level=2)
    doc.add_paragraph(f'Méthode utilisée : {method}')
    doc.add_paragraph(f'Statistique : {stat_value if stat_value else "N/A"}')
    doc.add_paragraph(f'Degrés de liberté : {dof_val if dof_val else "N/A"}')
    doc.add_paragraph(f'p-valeur : {p_val:.4f}')
    doc.add_paragraph(f'Conclusion : {"Pas d\'association significative" if p_val > 0.05 else "Association significative"}')
    
    output = BytesIO()
    doc.save(output)
    return output.getvalue()

def generate_quant_qual_report(df, quant_var, qual_var, norm_result, var_result, test_method, stat_val, p_val):
    doc = Document()
    doc.add_heading('Rapport d\'Analyse de Liaison : Quantitative vs Qualitative', 0)
    doc.add_heading(f'Variables analysées : {quant_var} vs {qual_var}', level=1)
    
    doc.add_heading('Résultats des Tests', level=2)
    doc.add_paragraph(f'Normalité ({quant_var}) : {norm_result} (p-valeur = {shapiro(df[quant_var])[1]:.4f})')
    doc.add_paragraph(f'Égalité des variances : {var_result}')
    doc.add_paragraph(f'Méthode utilisée : {test_method}')
    doc.add_paragraph(f'Statistique : {stat_val:.4f}')
    doc.add_paragraph(f'p-valeur : {p_val:.4f}')
    doc.add_paragraph(f'Conclusion : {"Pas d\'association significative" if p_val > 0.05 else "Association significative"}')
    
    output = BytesIO()
    doc.save(output)
    return output.getvalue()

# --- Sidebar ---
with st.sidebar:
    st.header("📁 Import des Données")
    use_demo = st.checkbox("Utiliser le jeu de données démo (Iris)", value=False)
    uploaded_file = None if use_demo else st.file_uploader("Choisir un fichier (CSV ou Excel)", type=["csv", "xlsx", "xls"])
    
    st.markdown("---")
    st.header("📊 Type d'Analyse")
    analysis_type = st.radio("Sélectionner le type d'analyse :", [
        "Quantitative vs Quantitative",
        "Qualitative vs Qualitative",
        "Quantitative vs Qualitative"
    ])
    
    df = None
    if use_demo:
        df = load_demo_data()
        st.success("Jeu de données démo (Iris) chargé avec succès")
    elif uploaded_file:
        try:
            if uploaded_file.name.endswith('.csv'):
                df = pd.read_csv(uploaded_file, sep=st.selectbox("Séparateur CSV", [",", ";", "\t"], index=1))
            else:
                df = pd.read_excel(uploaded_file)
            st.success(f"Fichier chargé : **{uploaded_file.name}**")
            st.write(f"🔍 **{len(df)}** observations, **{len(df.columns)}** variables")
            st.markdown(f"<p class='uploaded-file'>Fichier chargé : {uploaded_file.name}</p>", unsafe_allow_html=True)
        except Exception as e:
            st.error(f"Erreur lors de la lecture du fichier : {str(e)}")
            st.stop()
    
    if df is not None:
        numeric_cols = df.select_dtypes(include=['float64', 'int64']).columns.tolist()
        qualitative_cols = df.select_dtypes(include=['object', 'category']).columns.tolist()
        qualitative_cols += [col for col in df.columns if df[col].nunique() <= 10 and col not in qualitative_cols]
        
        if analysis_type == "Quantitative vs Quantitative":
            if len(numeric_cols) < 2:
                st.error("Au moins deux variables quantitatives sont requises.")
                st.stop()
            var1 = st.selectbox("Première variable quantitative", numeric_cols, key="quant1")
            var2 = st.selectbox("Seconde variable quantitative", [x for x in numeric_cols if x != var1], key="quant2")
            color_theme = st.selectbox("Thème de couleur", ["viridis", "plasma", "inferno", "magma"], key="color_theme")
        
        elif analysis_type == "Qualitative vs Qualitative":
            if len(qualitative_cols) < 2:
                st.error("Au moins deux variables qualitatives sont requises.")
                st.stop()
            var1 = st.selectbox("Première variable qualitative", qualitative_cols, key="qual1")
            var2 = st.selectbox("Seconde variable qualitative", [x for x in qualitative_cols if x != var1], key="qual2")
        
        else:  # Quantitative vs Qualitative
            if not numeric_cols or not qualitative_cols:
                st.error("Au moins une variable quantitative et une qualitative sont requises.")
                st.stop()
            var1 = st.selectbox("Variable quantitative", numeric_cols, key="quant_qual1")
            var2 = st.selectbox("Variable qualitative", qualitative_cols, key="quant_qual2")
    
    st.markdown("---")
    st.markdown("""
    <div class="author-info">
        <h4>🧾 À propos de l'auteur</h4>
        <p><b>Nom:</b> N'dri</p>
        <p><b>Prénom:</b> Abo Onesime</p>
        <p><b>Rôle:</b> Data Analyst / Scientist</p>
        <p><b>Téléphone:</b> 07-68-05-98-87 / 01-01-75-11-81</p>
        <p><b>Email:</b> <a href="mailto:ndriablatie123@gmail.com" style="color:#00ff88;">ndriablatie123@gmail.com</a></p>
        <p><b>LinkedIn:</b> <a href="https://www.linkedin.com/in/abo-onesime-n-dri-54a537200/" target="_blank" style="color:#00ff88;">Profil LinkedIn</a></p>
        <p><b>GitHub:</b> <a href="https://github.com/Aboonesime" target="_blank" style="color:#00ff88;">Mon GitHub</a></p>
    </div>
    """, unsafe_allow_html=True)

# --- Main Content ---
if df is not None and 'var1' in locals() and 'var2' in locals():
    st.header(f"Analyse : {analysis_type}")
    st.subheader("Aperçu des Données")
    st.dataframe(df.head(), use_container_width=True)
    
    if analysis_type == "Quantitative vs Quantitative":
        df_filtered = df[[var1, var2]].dropna()
        if len(df_filtered) < 3:
            st.warning("Pas assez de données valides après suppression des valeurs manquantes.")
            st.stop()
        
        # Correlation Matrix
        st.subheader("Matrice de Corrélation")
        corr_matrix = df_filtered.corr()
        fig_heatmap, ax_heatmap = plt.subplots(figsize=(8, 6))
        sns.heatmap(corr_matrix, annot=True, fmt=".4f", cmap=color_theme, linewidths=0.5, ax=ax_heatmap)
        ax_heatmap.set_title(f"Corrélation entre {var1} et {var2}", fontsize=16)
        st.pyplot(fig_heatmap)
        plt.close(fig_heatmap)
        
        # Normality Test
        st.subheader("Vérification de la Normalité")
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown(f"**{var1}**")
            fig1, ax1 = plt.subplots(figsize=(8, 5))
            sns.histplot(df_filtered[var1], kde=True, ax=ax1, color='#00ffff', stat="density")
            ax1.set_title(f"Distribution de {var1}")
            st.pyplot(fig1)
            plt.close(fig1)
            if len(df_filtered[var1]) >= 3 and df_filtered[var1].nunique() >= 3:
                stat1, p1 = shapiro(df_filtered[var1])
                st.write(f"Shapiro-Wilk : p-valeur = {p1:.4f}")
                normal_var1 = p1 > 0.05
                if normal_var1:
                    st.success("Distribution normale (p > 0.05)")
                else:
                    st.error("Distribution non normale (p ≤ 0.05)")
            else:
                st.warning("Pas assez de données pour le test de normalité.")
                normal_var1 = False
        
        with col2:
            st.markdown(f"**{var2}**")
            fig2, ax2 = plt.subplots(figsize=(8, 5))
            sns.histplot(df_filtered[var2], kde=True, ax=ax2, color='#00ff88', stat="density")
            ax2.set_title(f"Distribution de {var2}")
            st.pyplot(fig2)
            plt.close(fig2)
            if len(df_filtered[var2]) >= 3 and df_filtered[var2].nunique() >= 3:
                stat2, p2 = shapiro(df_filtered[var2])
                st.write(f"Shapiro-Wilk : p-valeur = {p2:.4f}")
                normal_var2 = p2 > 0.05
                if normal_var2:
                    st.success("Distribution normale (p > 0.05)")
                else:
                    st.error("Distribution non normale (p ≤ 0.05)")
            else:
                st.warning("Pas assez de données pour le test de normalité.")
                normal_var2 = False
        
        # Scatter Plot and Regression
        st.subheader("Nuage de Points et Régression")
        fig_scatter, ax_scatter = plt.subplots(figsize=(10, 6))
        sns.scatterplot(data=df_filtered, x=var1, y=var2, ax=ax_scatter, color='#00ffff', s=80, alpha=0.7)
        
        X_filtered = df_filtered[[var1]]
        y_filtered = df_filtered[var2]
        if len(X_filtered) >= 2 and X_filtered[var1].nunique() > 1:
            try:
                model = LinearRegression().fit(X_filtered, y_filtered)
                x_range = np.linspace(df_filtered[var1].min(), df_filtered[var1].max(), 100).reshape(-1, 1)
                y_pred = model.predict(x_range)
                ax_scatter.plot(x_range, y_pred, color='#00ff88', linewidth=2)
                ax_scatter.text(0.05, 0.95, f'y = {model.coef_[0]:.4f}x + {model.intercept_:.4f}',
                                transform=ax_scatter.transAxes, fontsize=12,
                                bbox=dict(facecolor='#1a1a2e', alpha=0.8, edgecolor='#00ffff'))
            except Exception as e:
                st.warning(f"Erreur lors de la régression : {str(e)}")
        
        ax_scatter.set_title(f"Nuage de points entre {var1} et {var2}")
        st.pyplot(fig_scatter)
        plt.close(fig_scatter)
        
        # Correlation Test
        st.subheader("Test de Corrélation")
        if normal_var1 and normal_var2:
            corr, p_value = pearsonr(df_filtered[var1], df_filtered[var2])
            coef_symbol = "r"
            test_name = "Pearson"
        else:
            corr, p_value = spearmanr(df_filtered[var1], df_filtered[var2])
            coef_symbol = "rho"
            test_name = "Spearman"
        
        value_box(f"Coefficient ({coef_symbol})", f"{corr:.4f}", "primary", "🧮")
        value_box("p-valeur", f"{p_value:.4f}", "success" if p_value > 0.05 else "info", "✅")
        st.write(f"Méthode utilisée : {test_name}")
        if p_value > 0.05:
            st.success("Pas d'association significative (p > 0.05)")
        else:
            st.info("Association significative (p ≤ 0.05)")
            col1, col2 = st.columns(2)
            with col1:
                st.write(f"Test positif ({coef_symbol} > 0) : p-valeur = {p_value/2:.4f}")
                if corr > 0 and p_value/2 < 0.05:
                    st.write("Corrélation positive significative.")
                else:
                    st.write("Pas de corrélation positive significative.")
            with col2:
                st.write(f"Test négatif ({coef_symbol} < 0) : p-valeur = {p_value/2:.4f}")
                if corr < 0 and p_value/2 < 0.05:
                    st.write("Corrélation négative significative.")
                else:
                    st.write("Pas de corrélation négative significative.")
        
        # Report Download
        report = generate_quant_quant_report(df_filtered, var1, var2, corr, p_value, coef_symbol, normal_var1, normal_var2)
        st.download_button(
            label="Télécharger le Rapport (Word)",
            data=report,
            file_name=f"rapport_quant_quant_{var1}_{var2}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    elif analysis_type == "Qualitative vs Qualitative":
        df_filtered = df[[var1, var2]].dropna()
        if len(df_filtered) < 3:
            st.warning("Pas assez de données valides après suppression des valeurs manquantes.")
            st.stop()
        
        # Contingency Table
        st.subheader("Tableau de Contingence")
        tab = pd.crosstab(df_filtered[var1], df_filtered[var2])
        st.dataframe(tab, use_container_width=True)
        
        # Cochran Conditions
        chi2, p, dof, expected = chi2_contingency(tab)
        expected_df = pd.DataFrame(expected, index=tab.index, columns=tab.columns).round(2)
        total_cells = expected.size
        cells_above_5 = (expected > 5).sum()
        percentage_above_5 = (cells_above_5 / total_cells) * 100
        cochran_ok = percentage_above_5 >= 80
        
        if cochran_ok:
            test_used = "Chi-Square"
            test_statistic = chi2
            test_df = dof
            p_value = p
        else:
            try:
                if tab.shape[0] > 2 or tab.shape[1] > 2 or tab.sum().sum() > 100:
                    _, p_value = fisher_exact(tab, simulate_p_value=True)
                    test_used = "Fisher Exact (Monte Carlo)"
                else:
                    _, p_value = fisher_exact(tab)
                    test_used = "Fisher Exact"
                test_statistic = None
                test_df = None
            except:
                test_used = "Chi-Square (Fisher failed)"
                test_statistic = chi2
                test_df = dof
                p_value = p
        
        st.subheader("Vérification des Conditions de Cochran")
        col1, col2, col3 = st.columns(3)
        with col1:
            value_box("Total Cellules", total_cells, "primary", "🔢")
        with col2:
            value_box("Cellules > 5", cells_above_5, "success", "✅")
        with col3:
            value_box("% Cellules > 5", f"{percentage_above_5:.1f}%", "success" if cochran_ok else "danger", "✅" if cochran_ok else "⚠️")
        
        if cochran_ok:
            st.success("Conditions de Cochran satisfaites.")
        else:
            st.warning("Conditions de Cochran non satisfaites. Test alternatif utilisé.")
        
        if (expected <= 5).any():
            st.subheader("Cellules Problématiques")
            problematic_df = expected_df.where(expected <= 5).stack().reset_index()
            problematic_df.columns = [var1, var2, "Effectif Théorique"]
            st.dataframe(problematic_df.dropna(), use_container_width=True)
        
        # Visualization
        st.subheader("Visualisation")
        plot_data = tab.reset_index().melt(id_vars=var1, var_name=var2, value_name="Fréquence")
        fig = px.bar(plot_data, x=var1, y="Fréquence", color=var2, barmode="group",
                     title=f"Relation entre {var1} et {var2}",
                     color_discrete_sequence=['#00ffff', '#00ff88', '#ffaa00', '#ff4444'],
                     template=plotly_dark_template)
        st.plotly_chart(fig, use_container_width=True)
        
        # Test Results
        st.subheader("Résultats du Test")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            value_box("Test", test_used, "primary", "🧮")
        with col2:
            value_box("Statistique", f"{test_statistic:.3f}" if test_statistic else "N/A", "primary", "📊")
        with col3:
            value_box("DDL", test_df if test_df else "N/A", "primary", "📏")
        with col4:
            value_box("p-valeur", f"{p_value:.4f}", "success" if p_value > 0.05 else "info", "✅")
        
        if p_value > 0.05:
            st.success("Pas d'association significative (p > 0.05)")
        else:
            st.info("Association significative (p ≤ 0.05)")
        
        # Report Download
        report = generate_qual_qual_report(tab, test_used, test_statistic, test_df, p_value)
        st.download_button(
            label="Télécharger le Rapport (Word)",
            data=report,
            file_name=f"rapport_qual_qual_{var1}_{var2}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    else:  # Quantitative vs Qualitative
        df_filtered = df[[var1, var2]].dropna()
        if len(df_filtered) < 3:
            st.warning("Pas assez de données valides après suppression des valeurs manquantes.")
            st.stop()
        
        # Normality Test
        st.subheader("Test de Normalité (Shapiro-Wilk)")
        shapiro_test = shapiro(df_filtered[var1])
        st.write(f"Statistique : {shapiro_test[0]:.4f}, p-valeur : {shapiro_test[1]:.4f}")
        normality_result = "Normale" if shapiro_test[1] > 0.05 else "Non normale"
        if normality_result == "Normale":
            st.success("Distribution normale (p > 0.05)")
        else:
            st.error("Distribution non normale (p ≤ 0.05)")
        
        # Variance Test
        st.subheader("Test d'Égalité des Variances")
        groups = [group[var1].values for name, group in df_filtered.groupby(var2)]
        if normality_result == "Normale":
            bartlett_test = bartlett(*groups)
            st.write(f"Bartlett : Statistique = {bartlett_test[0]:.4f}, p-valeur = {bartlett_test[1]:.4f}")
            variance_result = "Égales" if bartlett_test[1] > 0.05 else "Inégales"
            if variance_result == "Égales":
                st.success("Variances égales (p > 0.05)")
            else:
                st.error("Variances inégales (p ≤ 0.05)")
        else:
            levene_test = levene(*groups)
            st.write(f"Levene : Statistique = {levene_test[0]:.4f}, p-valeur = {levene_test[1]:.4f}")
            variance_result = "Égales" if levene_test[1] > 0.05 else "Inégales"
            if variance_result == "Égales":
                st.success("Variances égales (p > 0.05)")
            else:
                st.error("Variances inégales (p ≤ 0.05)")
        
        # Association Test
        st.subheader("Test de Liaison")
        if normality_result == "Normale" and variance_result == "Égales":
            anova_test = f_oneway(*groups)
            test_name = "ANOVA"
            test_statistic = anova_test[0]
            p_value = anova_test[1]
        else:
            kruskal_test = kruskal(*groups)
            test_name = "Kruskal-Wallis"
            test_statistic = kruskal_test[0]
            p_value = kruskal_test[1]
        
        st.write(f"{test_name} : Statistique = {test_statistic:.4f}, p-valeur = {p_value:.4f}")
        if p_value > 0.05:
            st.success("Pas d'association significative (p > 0.05)")
        else:
            st.info("Association significative (p ≤ 0.05)")
        
        # Visualization
        st.subheader("Visualisation")
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(15, 6))
        sns.boxplot(data=df_filtered, x=var2, y=var1, ax=ax1, color='#00ffff')
        ax1.set_title(f"Boxplot de {var1} par {var2}")
        sns.violinplot(data=df_filtered, x=var2, y=var1, ax=ax2, color='#00ffff')
        ax2.set_title(f"Violin Plot de {var1} par {var2}")
        st.pyplot(fig)
        plt.close(fig)
        
        # QQ Plot
        st.subheader("QQ-Plot")
        fig2, ax = plt.subplots(figsize=(8, 6))
        probplot(df_filtered[var1], dist="norm", plot=ax)
        ax.get_lines()[0].set_markerfacecolor('#00ffff')
        ax.get_lines()[0].set_markeredgecolor('#00ffff')
        ax.get_lines()[1].set_color('red')
        ax.set_title(f"QQ-Plot de {var1}")
        st.pyplot(fig2)
        plt.close(fig2)
        
        # Python Code
        st.subheader("Code Python")
        code = f"""import pandas as pd
from scipy import stats
import matplotlib.pyplot as plt

df = pd.read_csv('your_data.csv')  # Remplacez par votre fichier
var1 = '{var1}'
var2 = '{var2}'

# Normalité
shapiro_result = stats.shapiro(df[var1])
print(f"Shapiro-Wilk: Statistique={{shapiro_result[0]:.3f}}, p-valeur={{shapiro_result[1]:.3f}}")
normality_result = "Normale" if shapiro_result[1] > 0.05 else "Non normale"

# Variances
groups = [group[var1].values for name, group in df.groupby(var2)]
if normality_result == "Normale":
    bartlett_result = stats.bartlett(*groups)
    print(f"Bartlett: Statistique={{bartlett_result[0]:.3f}}, p-valeur={{bartlett_result[1]:.3f}}")
    variance_result = "Égales" if bartlett_result[1] > 0.05 else "Inégales"
else:
    levene_result = stats.levene(*groups)
    print(f"Levene: Statistique={{levene_result[0]:.3f}}, p-valeur={{levene_result[1]:.3f}}")
    variance_result = "Égales" if levene_result[1] > 0.05 else "Inégales"

# Test de liaison
if normality_result == "Normale" and variance_result == "Égales":
    anova_result = stats.f_oneway(*groups)
    print(f"ANOVA: Statistique={{anova_result[0]:.3f}}, p-valeur={{anova_result[1]:.3f}}")
else:
    kruskal_result = stats.kruskal(*groups)
    print(f"Kruskal-Wallis: Statistique={{kruskal_result[0]:.3f}}, p-valeur={{kruskal_result[1]:.3f}}")

# Visualisation
plt.figure(figsize=(10, 6))
df.boxplot(column=var1, by=var2)
plt.title(f"Distribution de {var1} par {var2}")
plt.suptitle('')
plt.show()
"""
        st.code(code, language='python')
        st.download_button(
            label="Télécharger le Code Python",
            data=code,
            file_name="statistical_test.py",
            mime="text/plain"
        )
        
        # Report Download
        report = generate_quant_qual_report(df_filtered, var1, var2, normality_result, variance_result, test_name, test_statistic, p_value)
        st.download_button(
            label="Télécharger le Rapport (Word)",
            data=report,
            file_name=f"rapport_quant_qual_{var1}_{var2}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

else:
    st.info("Veuillez uploader un fichier CSV/Excel ou utiliser le jeu de données démo, puis sélectionner un type d'analyse et des variables.")
    st.markdown("""
    ### Instructions :
    1. Choisissez un fichier CSV/Excel ou utilisez le jeu de données démo (Iris).
    2. Sélectionnez le type d'analyse :
       - Quantitative vs Quantitative : Tests de corrélation (Pearson/Spearman).
       - Qualitative vs Qualitative : Tests d'indépendance (Chi-Square/Fisher).
       - Quantitative vs Qualitative : Tests de comparaison (ANOVA/Kruskal-Wallis).
    3. Sélectionnez les variables appropriées.
    4. Explorez les résultats, visualisations, et téléchargez les rapports ou le code Python.
    """, unsafe_allow_html=True)

# --- Footer ---
st.markdown("---")
st.markdown("""
<div class="footer">
    Application d'Analyse de Liaison Statistique • Créée avec Streamlit & Plotly • © 2025 Abo Onesime
</div>
""", unsafe_allow_html=True)